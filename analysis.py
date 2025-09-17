import pandas as pd
import numpy as np
from scipy.stats import ttest_ind, mannwhitneyu, shapiro
from docx import Document
from docx.shared import Inches
import warnings
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Rectangle, Wedge
from matplotlib.collections import PatchCollection

# === FONCTIONS VISUALISATION ===
def create_boxplots(data, vec_columns, group_col, group_labels, title, results_df=None):
    """Crée des boîtes à moustaches pour les valeurs VEC selon les groupes."""
    plt.figure(figsize=(15, 10))
    for i, col in enumerate(vec_columns, 1):
        plt.subplot(3, 4, i)
        ax = plt.gca()
        sns.boxplot(x=group_col, y=col, data=data, order=group_labels)
        plt.title(col.replace('VEC_', ''))
        plt.xticks(rotation=45)
        
        # Ajouter les étoiles pour les différences significatives
        if results_df is not None:
            # Trouver les résultats significatifs pour cette variable
            sig_results = results_df[
                (results_df['Variable'] == col) & 
                (results_df['Signif'] == '**')
            ]
            
            if not sig_results.empty:
                y_max = data[col].max()
                y_range = data[col].max() - data[col].min()
                y_top = y_max + 0.1 * y_range  # Marge pour les annotations
                
                # Espacer les lignes de significativité si plusieurs comparaisons
                for idx, row in enumerate(sig_results.iterrows()):
                    row = row[1]  # Accéder aux données de la ligne
                    # Trouver les indices des groupes comparés
                    comp_groups = row['Comparaison'].split(' vs ')
                    try:
                        idx1 = group_labels.index(comp_groups[0])
                        idx2 = group_labels.index(comp_groups[1])
                        
                        # Calculer la position des étoiles (plus haut pour chaque comparaison)
                        y_pos = y_top + (idx * 0.05 * y_range)
                        
                        # Tracer la ligne de connexion
                        x1, x2 = idx1, idx2
                        plt.plot([x1, x2], [y_pos, y_pos], 'k-', linewidth=1)
                        
                        # Ajouter les étoiles
                        plt.text((x1 + x2) / 2, y_pos + 0.01 * y_range, '**', 
                               ha='center', va='bottom')
                    except ValueError:
                        continue
                
                # Ajuster les limites pour s'assurer que tout est visible
                current_ymax = ax.get_ylim()[1]
                if sig_results.shape[0] > 0:  # S'il y a des résultats significatifs
                    new_ymax = y_top + (sig_results.shape[0] * 0.05 + 0.05) * y_range
                    ax.set_ylim(ax.get_ylim()[0], max(current_ymax, new_ymax))
                
    plt.tight_layout()
    return plt.gcf()

def create_bullseye(data, prefix, title, include_apex=False):
    """Crée un diagramme bullseye pour les valeurs VEC moyennes."""
    # Configuration
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
    
    num_theta = 4  # 4 segments par anneau
    num_radii = 3  # Toujours inclure l'apex visuellement
    thetas = np.linspace(0, 2 * np.pi, num_theta + 1)
    
    # Labels pour les axes
    radii_labels = ['Apex', 'Mid', 'Base']
    theta_labels = ['Latéral', 'Antérieur', 'Septal', 'Inférieur']
    
    # Configuration des axes
    ax.set_theta_offset(np.pi / 4)  # Fait pivoter pour que 0° soit à droite
    ax.set_theta_direction(-1)  # Met en sens horaire
    ax.set_xticks(np.linspace(0, 2 * np.pi, num_theta, endpoint=False))
    ax.set_xticklabels(theta_labels, fontsize=12)
    # Ajuster les ticks en fonction du nombre d'anneaux
    if include_apex:
        ax.set_yticks([0.5, 1.5, 2.5])
    else:
        ax.set_yticks([0.5, 1.5, 2.5])
    ax.set_yticklabels(radii_labels, fontsize=12)
    ax.set_ylim(0, num_radii)
    ax.tick_params(pad=20)
    
    # Dessiner l'apex en grisé si non inclus
    if not include_apex:
        for t_idx in range(num_theta):
            ax.bar(x=thetas[t_idx], height=1, width=(2*np.pi/num_theta),
                  bottom=0, color='lightgray', edgecolor='white',
                  linewidth=2, alpha=0.5)
            # Ajouter le texte "Non utilisé"
            angle = thetas[t_idx] + np.pi/num_theta
            ax.text(angle, 0.5, "Non utilisé", ha='center', va='center', 
                   fontsize=8, color='gray', alpha=0.7)
    
    # Mapping des segments aux positions
    base_segments = {
        'base_lat': (1, 0), 'base_ante': (1, 1), 'base_sept': (1, 2), 'base_inf': (1, 3)
    }
    mid_segments = {
        'mid_lat': (0, 0), 'mid_ante': (0, 1), 'mid_sept': (0, 2), 'mid_inf': (0, 3)
    }
    apex_segments = {
        'apex_lat': (0, 0), 'apex_ante': (0, 1), 'apex_sept': (0, 2), 'apex_inf': (0, 3)
    }
    
    if include_apex:
        segment_mapping = {**apex_segments, **{k: (v[0]+1, v[1]) for k, v in mid_segments.items()},
                         **{k: (v[0]+2, v[1]) for k, v in base_segments.items()}}
    else:
        segment_mapping = {**{k: (v[0]+1, v[1]) for k, v in mid_segments.items()},
                         **{k: (v[0]+2, v[1]) for k, v in base_segments.items()}}
    
    # Calcul des valeurs min/max pour l'échelle de couleur
    values = []
    levels = ['base', 'mid', 'apex'] if include_apex else ['base', 'mid']
    for level in levels:
        for pos in ['lat', 'ante', 'sept', 'inf']:
            col_name = f"VEC_{level}_{pos}"
            if col_name in data.columns:
                mean_val = data[col_name].mean()
                values.append(mean_val)
    
    vmin, vmax = min(values), max(values)
    
    # Création de la colormap avec couleurs inversées
    cmap = plt.cm.viridis_r  # _r pour inverser l'échelle de couleur
    norm = plt.Normalize(vmin=vmin, vmax=vmax)
    
    # Dessin des segments avec les valeurs moyennes et écarts-types
    for segment_base, (r_idx, t_idx) in segment_mapping.items():
        col_name = f"VEC_{segment_base}"
        if col_name in data.columns:
            mean_val = data[col_name].mean()
            std_val = data[col_name].std()
            color = cmap(norm(mean_val))
            
            # Dessiner le segment
            ax.bar(x=thetas[t_idx], height=1, width=(2*np.pi/num_theta),
                  bottom=r_idx, color=color, edgecolor='white',
                  linewidth=2, alpha=0.9)
            
            # Ajouter le texte avec moyenne ± écart-type
            angle = thetas[t_idx] + np.pi/num_theta
            r = r_idx + 0.5
            ha = 'center'
            text = f"{mean_val:.3f}\n±{std_val:.3f}"
            ax.text(angle, r, text, ha=ha, va='center', fontsize=8)
    
    # Ajouter la barre de couleur
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = plt.colorbar(sm, ax=ax, pad=0.1, shrink=0.7)
    cbar.set_label('Valeur Moyenne VEC', rotation=270, labelpad=20, fontsize=12)
    
    ax.set_title(title, fontsize=14, pad=45)
    plt.tight_layout()
    
    return fig

# === PARAMÈTRES ===
file_path = "CRF_Fabry_unique_modifCF.xlsx"  # mets le chemin complet si besoin
sheet_name = "Données Fabry suivis"  # adapte si le nom diffère

# === CHARGEMENT ===
df = pd.read_excel(file_path, sheet_name=sheet_name)

# Harmonisation de la colonne Sexe
df["Sexe"] = df["Sexe"].replace({
    "M": "H",  # Convertir M en H pour Homme
    "Hommes": "H",
    "F": "F",  # Garder F pour Femme
    "Femmes": "F"
}).astype(str)

# Nettoyage : garder uniquement les lignes avec H ou F
df = df[df["Sexe"].isin(["H", "F"])]

# Colonnes d'intérêt : toutes les colonnes VEC
vec_cols = [c for c in df.columns if "VEC" in c]

# Harmonisation des colonnes catégorielles
def harmonize_binary_column(series):
    """Convert mixed data types to binary (0/1) values."""
    # Convert to string first to handle all types uniformly
    series = series.astype(str)
    
    # Handle categorical values
    series = series.replace({
        "Oui": "1", "Non": "0"
    })
    
    # Try to convert to float first to handle decimal numbers
    numeric_series = pd.to_numeric(series, errors='coerce')
    # Convert any numeric value > 0 to 1, everything else to 0
    numeric_series = (numeric_series > 0).astype(int)
    
    return numeric_series

# Apply harmonization to both columns
df["fibrose"] = harmonize_binary_column(df["fibrose"])
df["HVG"] = harmonize_binary_column(df["HVG"])

# === FONCTION COMPARAISON ===
def compare_groups(data, col, cond1, cond2, label1, label2):
    """Compare une variable continue entre deux groupes (cond1 vs cond2)."""
    # Convert to numeric, dropping any non-numeric values
    g1 = pd.to_numeric(data.loc[cond1, col], errors='coerce').dropna()
    g2 = pd.to_numeric(data.loc[cond2, col], errors='coerce').dropna()
    n1, n2 = len(g1), len(g2)
    if n1 < 3 or n2 < 3:
        return None

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        p1, p2 = shapiro(g1)[1], shapiro(g2)[1]

    if p1 > 0.05 and p2 > 0.05:
        stat, p = ttest_ind(g1, g2, equal_var=False)
        test = "t-test"
    else:
        stat, p = mannwhitneyu(g1, g2)
        test = "Mann-Whitney"

    signif = "**" if p < 0.05 else ""
    return {
        "Variable": col,
        "Comparaison": f"{label1} vs {label2}",
        "N": n1 + n2,
        "Moyenne ± ET (grp1)": f"{g1.mean():.3f} ± {g1.std():.3f}",
        "Moyenne ± ET (grp2)": f"{g2.mean():.3f} ± {g2.std():.3f}",
        "Test": test,
        "p-value": round(p, 4),
        "Signif": signif
    }

# === LISTE DES COMPARAISONS ===
comparaisons = [
    # Homme vs Femme
    ("Sexe == 'H'", "Sexe == 'F'", "Hommes", "Femmes"),
    # Homme fibrose vs Homme sans fibrose
    ("Sexe == 'H' & fibrose == 1", "Sexe == 'H' & fibrose == 0", "H Fibrose", "H sans Fibrose"),
    # Femme fibrose vs Femme sans fibrose
    ("Sexe == 'F' & fibrose == 1", "Sexe == 'F' & fibrose == 0", "F Fibrose", "F sans Fibrose"),
    # Homme HVG vs Homme sans HVG
    ("Sexe == 'H' & HVG == 1", "Sexe == 'H' & HVG == 0", "H HVG", "H sans HVG"),
    # Femme HVG vs Femme sans HVG
    ("Sexe == 'F' & HVG == 1", "Sexe == 'F' & HVG == 0", "F HVG", "F sans HVG"),
    # H avec fibrose vs F avec fibrose
    ("Sexe == 'H' & fibrose == 1", "Sexe == 'F' & fibrose == 1", "H Fibrose", "F Fibrose"),
    # H sans fibrose vs F sans fibrose
    ("Sexe == 'H' & fibrose == 0", "Sexe == 'F' & fibrose == 0", "H sans Fibrose", "F sans Fibrose"),
]

results = []
for cond1, cond2, label1, label2 in comparaisons:
    for col in vec_cols:
        res = compare_groups(df, col, df.eval(cond1), df.eval(cond2), label1, label2)
        if res:
            results.append(res)

results_df = pd.DataFrame(results)

# === VISUALISATIONS ===
# 1. Boîtes à moustaches pour la comparaison Homme vs Femme
boxplot_sex = create_boxplots(df, vec_cols, 'Sexe', ['H', 'F'], 
                            'Comparaison VEC Homme vs Femme',
                            results_df[results_df['Comparaison'] == 'Hommes vs Femmes'])
boxplot_sex.savefig('VEC_boxplot_sex.png')

# 2. Boîtes à moustaches pour la comparaison avec/sans fibrose par sexe
df_fibrose = df.copy()
df_fibrose['Groupe'] = df_fibrose.apply(
    lambda x: f"{'H' if x['Sexe']=='H' else 'F'} {'Fibrose' if x['fibrose']==1 else 'sans Fibrose'}", 
    axis=1
)
# Filtrer les résultats pour la fibrose
fibrose_results = results_df[
    (results_df['Comparaison'].str.contains('Fibrose vs')) & 
    (results_df['Comparaison'].str.contains('sans Fibrose'))
]
boxplot_fibrose = create_boxplots(df_fibrose, vec_cols, 'Groupe', 
                                 ['H Fibrose', 'H sans Fibrose', 'F Fibrose', 'F sans Fibrose'],
                                 'Comparaison VEC avec/sans Fibrose par Sexe',
                                 fibrose_results)
boxplot_fibrose.savefig('VEC_boxplot_fibrose.png')

# 3. Boîtes à moustaches pour la comparaison avec/sans HVG par sexe
df_hvg = df.copy()
df_hvg['Groupe'] = df_hvg.apply(
    lambda x: f"{'H' if x['Sexe']=='H' else 'F'} {'HVG' if x['HVG']==1 else 'sans HVG'}", 
    axis=1
)
# Filtrer les résultats pour HVG
hvg_results = results_df[
    (results_df['Comparaison'].str.contains('HVG vs')) & 
    (results_df['Comparaison'].str.contains('sans HVG'))
]
boxplot_hvg = create_boxplots(df_hvg, vec_cols, 'Groupe', 
                            ['H HVG', 'H sans HVG', 'F HVG', 'F sans HVG'],
                            'Comparaison VEC avec/sans HVG par Sexe',
                            hvg_results)
boxplot_hvg.savefig('VEC_boxplot_hvg.png')

# === DIAGRAMMES BULLSEYE ===
# Comparaison globale Hommes vs Femmes (sans apex)
men_all = df[df['Sexe'] == 'H']
women_all = df[df['Sexe'] == 'F']

# Calcul de N pour chaque groupe
n_men = len(men_all)
n_women = len(women_all)

bullseye_men = create_bullseye(men_all, 'VEC', f'Bullseye VEC - Hommes (N={n_men})', include_apex=False)
bullseye_men.savefig('VEC_bullseye_H.png')
bullseye_women = create_bullseye(women_all, 'VEC', f'Bullseye VEC - Femmes (N={n_women})', include_apex=False)
bullseye_women.savefig('VEC_bullseye_F.png')

# On ne génère plus les autres bullseye car nous voulons juste la vue globale H vs F
men_fibrose = df[(df['Sexe'] == 'H') & (df['fibrose'] == 1)]
bullseye_men_fibrose = create_bullseye(men_fibrose, 'VEC', 'Bullseye VEC - Hommes avec Fibrose')
bullseye_men_fibrose.savefig('VEC_bullseye_H_fibrose.png')

# Hommes sans fibrose
men_no_fibrose = df[(df['Sexe'] == 'H') & (df['fibrose'] == 0)]
bullseye_men_no_fibrose = create_bullseye(men_no_fibrose, 'VEC', 'Bullseye VEC - Hommes sans Fibrose')
bullseye_men_no_fibrose.savefig('VEC_bullseye_H_sans_fibrose.png')

# Femmes avec fibrose
women_fibrose = df[(df['Sexe'] == 'F') & (df['fibrose'] == 1)]
bullseye_women_fibrose = create_bullseye(women_fibrose, 'VEC', 'Bullseye VEC - Femmes avec Fibrose')
bullseye_women_fibrose.savefig('VEC_bullseye_F_fibrose.png')

# Femmes sans fibrose
women_no_fibrose = df[(df['Sexe'] == 'F') & (df['fibrose'] == 0)]
bullseye_women_no_fibrose = create_bullseye(women_no_fibrose, 'VEC', 'Bullseye VEC - Femmes sans Fibrose')
bullseye_women_no_fibrose.savefig('VEC_bullseye_F_sans_fibrose.png')

# 3. Comparaison avec/sans HVG par sexe
# Hommes avec HVG
men_hvg = df[(df['Sexe'] == 'H') & (df['HVG'] == 1)]
bullseye_men_hvg = create_bullseye(men_hvg, 'VEC', 'Bullseye VEC - Hommes avec HVG')
bullseye_men_hvg.savefig('VEC_bullseye_H_hvg.png')

# Hommes sans HVG
men_no_hvg = df[(df['Sexe'] == 'H') & (df['HVG'] == 0)]
bullseye_men_no_hvg = create_bullseye(men_no_hvg, 'VEC', 'Bullseye VEC - Hommes sans HVG')
bullseye_men_no_hvg.savefig('VEC_bullseye_H_sans_hvg.png')

# Femmes avec HVG
women_hvg = df[(df['Sexe'] == 'F') & (df['HVG'] == 1)]
bullseye_women_hvg = create_bullseye(women_hvg, 'VEC', 'Bullseye VEC - Femmes avec HVG')
bullseye_women_hvg.savefig('VEC_bullseye_F_hvg.png')

# Femmes sans HVG
women_no_hvg = df[(df['Sexe'] == 'F') & (df['HVG'] == 0)]
bullseye_women_no_hvg = create_bullseye(women_no_hvg, 'VEC', 'Bullseye VEC - Femmes sans HVG')
bullseye_women_no_hvg.savefig('VEC_bullseye_F_sans_hvg.png')

plt.close('all')  # Fermer toutes les figures pour libérer la mémoire

# === EXPORT WORD ===
doc = Document()
doc.add_heading("Analyse comparative VEC Fabry", 0)
doc.add_paragraph("Toutes les comparaisons incluent le nombre de patients, "
                  "la moyenne ± ET, le test utilisé et l'indication de significativité (** si p < 0.05).")

table = doc.add_table(rows=1, cols=len(results_df.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(results_df.columns):
    hdr_cells[i].text = col

for _, row in results_df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Ajouter les images au document
doc.add_heading("Visualisations", level=1)

# Ajouter les boîtes à moustaches
doc.add_heading("Boîtes à moustaches", level=2)
doc.add_picture('VEC_boxplot_sex.png', width=Inches(6))
doc.add_paragraph("Figure 1: Comparaison des valeurs VEC entre Hommes et Femmes")
doc.add_picture('VEC_boxplot_fibrose.png', width=Inches(6))
doc.add_paragraph("Figure 2: Comparaison des valeurs VEC selon la présence de fibrose et le sexe")

# Ajouter les diagrammes Bullseye
doc.add_heading("Diagrammes Bullseye - Comparaison Hommes vs Femmes", level=2)
doc.add_picture('VEC_bullseye_H.png', width=Inches(4))
doc.add_paragraph("Figure 3: Distribution des VEC chez les Hommes (tous confondus)")
doc.add_picture('VEC_bullseye_F.png', width=Inches(4))
doc.add_paragraph("Figure 4: Distribution des VEC chez les Femmes (toutes confondues)")

doc.save("Analyse_VEC_Fabry_comparaisons.docx")
print("✅ Analyse terminée ! Résultats exportés dans : Analyse_VEC_Fabry_comparaisons.docx")